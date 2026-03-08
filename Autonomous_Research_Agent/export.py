import re
from docx import Document

def _add_paragraph_with_bold(doc, text, style=None):
    """
    自訂函數：將帶有 **粗體** 的 Markdown 單行文字，寫入 Word (Docx) 中的特定段落。
    """
    if style:
        p = doc.add_paragraph(style=style)
    else:
        p = doc.add_paragraph()
        
    # 利用正則表達式把文本切塊，偵測出包裹在 **...** 裡面的文字
    parts = re.split(r'(\*\*.*?\*\*)', text)
    for part in parts:
        if part.startswith('**') and part.endswith('**'):
            p.add_run(part[2:-2]).bold = True
        else:
            p.add_run(part)

def markdown_to_docx(md_text: str, output_path: str):
    """
    將 Gemini 回傳的標準 Markdown 文字內容解析為 Word 文件。
    支援標題 (H1, H2, H3)、無序列表、有序列表以及粗體格式。
    """
    doc = Document()
    
    lines = md_text.split('\n')
    for line in lines:
        stripped = line.strip()
        if not stripped:
            continue
            
        # 處理標題
        if stripped.startswith('# '):
            doc.add_heading(stripped[2:].strip(), level=1)
        elif stripped.startswith('## '):
            doc.add_heading(stripped[3:].strip(), level=2)
        elif stripped.startswith('### '):
            doc.add_heading(stripped[4:].strip(), level=3)
            
        # 處理無序列表
        elif stripped.startswith('- ') or stripped.startswith('* '):
            _add_paragraph_with_bold(doc, stripped[2:].strip(), style='List Bullet')
            
        # 處理有序列表，例如 "1. 項目"
        elif re.match(r'^\d+\.\s', stripped):
            text = re.sub(r'^\d+\.\s', '', stripped).strip()
            _add_paragraph_with_bold(doc, text, style='List Number')
            
        # 處理引言區塊
        elif stripped.startswith('> '):
            _add_paragraph_with_bold(doc, stripped[2:].strip(), style='Quote')
            
        # 處理一般文字
        else:
            _add_paragraph_with_bold(doc, stripped)
            
    doc.save(output_path)
    return output_path
